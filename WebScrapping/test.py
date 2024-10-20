import os
import requests
from bs4 import BeautifulSoup
from docx import Document

while True:
    url = input("Please enter the job listing URL (or type 'n' to exit): ")

    if url.lower() == 'n':  # Check if the user wants to exit
        print("Exiting the program.")
        break

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36'
    }

    response = requests.get(url, headers=headers)

    if response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')

        h1_element = soup.find('h1', class_='top-card-layout__title')
        company_link = soup.find('a', class_='topcard__org-name-link')
        location_span = soup.find('span', class_='topcard__flavor topcard__flavor--bullet')

        job_title = h1_element.get_text(strip=True) if h1_element else "N/A"
        company_name = company_link.get_text(strip=True) if company_link else "N/A"
        location = location_span.get_text(strip=True) if location_span else "N/A"

        file_name = 'job_details.docx'

        if os.path.exists(file_name):
            document = Document(file_name)
            table = document.tables[0]
            row_cells = table.add_row().cells
            row_cells[0].text = company_name + ", " + location
            row_cells[1].text = job_title
            
            print("Appended new job details to the existing document.")
          
        else:
            document = Document()
            table = document.add_table(rows=1, cols=2)
            row_cells = table.add_row().cells
            row_cells[0].text = company_name + ", " + location
            row_cells[1].text = job_title

            print("Created a new document with job details.")

        document.save(file_name)
    else:
        print(f"Failed to retrieve page. Status code: {response.status_code}")
