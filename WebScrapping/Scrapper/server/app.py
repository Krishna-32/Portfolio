from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
import os
import re
import requests
from bs4 import BeautifulSoup
from docx import Document
from pathlib import Path
import time

app = Flask(__name__)
CORS(app)

# Function to get downloads folder path
def get_downloads_folder():
    home = str(Path.home())  # Get the home directory
    downloads_path = os.path.join(home, 'Downloads')  # Append 'Downloads' to home directory
    return downloads_path

# Function to validate LinkedIn job URL
def is_valid_linkedin_url(url):
    pattern = r'https:\/\/www\.linkedin\.com\/jobs\/view\/\d+\/\?[^&]+'
    return re.match(pattern, url) is not None

# Function to request the job URL with retry on status code 429
def fetch_job_details_with_retry(url, headers, max_retries=5, delay=2):
    attempt = 0
    while attempt < max_retries:
        response = requests.get(url, headers=headers)

        if response.status_code == 200:
            return response  # Success, return the response
        
        elif response.status_code == 429:
            attempt += 1
            print(f"Received status code 429, retrying in {delay} seconds... (Attempt {attempt}/{max_retries})")
            time.sleep(delay)  # Wait before retrying
        
        else:
            return response  # For other errors, return the response immediately

    return None  # If max retries exceeded, return None


@app.route('/', methods=['POST'])
def add_job():
    data = request.json
    url = data.get('url')

    if not url:
        return jsonify({"error": "URL is required"}), 400
    
    # Validate the URL format
    if not is_valid_linkedin_url(url):
        return jsonify({"error": "Invalid LinkedIn job URL. Please provide a valid URL."}), 400

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/130.0.0.0 Safari/537.36'
    }

    response = fetch_job_details_with_retry(url, headers)

    if response and response.status_code == 200:
        soup = BeautifulSoup(response.content, 'html.parser')

        h1_element = soup.find('h1', class_='top-card-layout__title')
        company_link = soup.find('a', class_='topcard__org-name-link')
        location_span = soup.find('span', class_='topcard__flavor topcard__flavor--bullet')

        job_title = h1_element.get_text(strip=True) if h1_element else "N/A"
        company_name = company_link.get_text(strip=True) if company_link else "N/A"
        location = location_span.get_text(strip=True) if location_span else "N/A"

        downloads_folder = get_downloads_folder()
        file_name = os.path.join(downloads_folder, 'job_details.docx')  # Save in Downloads folder

        try:
            if os.path.exists(file_name):
                # Try to open the document to update it
                try:
                    document = Document(file_name)
                    table = document.tables[0]
                    row_cells = table.add_row().cells
                    row_cells[0].text = company_name + ", " + location
                    row_cells[1].text = job_title
                    document.save(file_name)
                except PermissionError:
                    return jsonify({"error": "Please close the file to continue updating."}), 400  # File is open
            else:
                # Create a new document
                document = Document()
                table = document.add_table(rows=1, cols=2)
                row_cells = table.add_row().cells
                row_cells[0].text = company_name + ", " + location
                row_cells[1].text = job_title
                document.save(file_name)

            return jsonify({
                "message": "Job details added successfully",
                "file_path": file_name,
                "job_title": job_title,
                "company_name": company_name,
                "location": location
            }), 200

        except Exception as e:
            return jsonify({"error": str(e)}), 500

    elif response:
        return jsonify({"error": f"Failed to retrieve page. Status code: {response.status_code}"}), 400
    else:
        return jsonify({"error": "Failed to retrieve page after multiple retries. Try Again."}), 500


@app.route('/')
def a():
    return render_template('index.html')


if __name__ == '__main__':
    app.run(debug=True)
